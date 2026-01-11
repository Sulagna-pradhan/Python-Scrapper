import requests
from bs4 import BeautifulSoup
from fpdf import FPDF
from docx import Document
import os

# Ensure these paths are correct for your environment
FONT_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
FONT_MONO = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"

def fetch_article(url):
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
    res = requests.get(url, headers=headers)
    res.raise_for_status()

    soup = BeautifulSoup(res.text, "html.parser")
    title_tag = soup.find("h1")
    title = title_tag.get_text(strip=True) if title_tag else "GeeksforGeeks Article"
    
    article = soup.find("div", class_="article-body") or soup.find("div", class_="text")

    return title, article

def parse_blocks(article):
    blocks = []
    if not article:
        return blocks

    # We use find_all but we will skip tags that are inside a table we already processed
    processed_tables = []

    for tag in article.find_all(["h2", "h3", "p", "pre", "ul", "table"], recursive=True):
        
        # FIX: Check if this tag is inside a table we've already parsed to avoid duplication
        if any(parent.name == 'table' for parent in tag.parents):
            continue

        if tag.name in ["h2", "h3"]:
            blocks.append(("heading", tag.get_text(strip=True)))
        
        elif tag.name == "p":
            text = tag.get_text(" ", strip=True)
            if text:
                blocks.append(("paragraph", text))
        
        elif tag.name == "pre":
            code = tag.get_text()
            if code.strip():
                blocks.append(("code", code))
        
        elif tag.name == "ul":
            # Avoid list items being repeated if they are inside tables
            items = [li.get_text(" ", strip=True) for li in tag.find_all("li", recursive=False)]
            for item in items:
                blocks.append(("list", item))
        
        elif tag.name == "table":
            table_data = []
            rows = tag.find_all("tr")
            for row in rows:
                cols = row.find_all(["td", "th"])
                col_text = [ele.get_text(strip=True) for ele in cols]
                if col_text:
                    table_data.append(col_text)
            if table_data:
                blocks.append(("table", table_data))

    return blocks

def save_pdf(title, blocks, filename):
    # Using fpdf2 (pip install fpdf2)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    try:
        pdf.add_font("DejaVu", "", FONT_REGULAR)
        pdf.add_font("DejaVu", "B", FONT_BOLD)
        pdf.add_font("DejaVuMono", "", FONT_MONO)
        font_style, mono_style = "DejaVu", "DejaVuMono"
    except:
        font_style, mono_style = "Helvetica", "Courier"

    # Title
    pdf.set_font(font_style, "B", 18)
    pdf.multi_cell(0, 10, title)
    pdf.ln(4)

    usable_width = pdf.w - 2 * pdf.l_margin

    for kind, content in blocks:
        if kind == "heading":
            pdf.set_font(font_style, "B", 14); pdf.ln(4)
            pdf.multi_cell(0, 9, content); pdf.ln(2)
        elif kind == "paragraph":
            pdf.set_font(font_style, "", 11)
            pdf.multi_cell(0, 7, content); pdf.ln(2)
        elif kind == "list":
            pdf.set_font(font_style, "", 11); pdf.set_x(pdf.l_margin + 10)
            pdf.multi_cell(usable_width - 10, 7, f"• {content}"); pdf.ln(1)
        elif kind == "code":
            pdf.set_font(mono_style, "", 9); pdf.set_fill_color(240, 240, 240)
            pdf.multi_cell(0, 5, content, fill=True); pdf.ln(3)
        elif kind == "table":
            pdf.set_font(font_style, "", 9)
            # This ensures table content doesn't repeat and stays formatted
            with pdf.table(borders_layout="SINGLE_TOP_LINE", cell_fill_color=(245, 245, 245), cell_fill_mode="ROWS", line_height=7) as table:
                for data_row in content:
                    row = table.row()
                    for datum in data_row:
                        row.cell(datum)
            pdf.ln(4)

    pdf.output(filename)

def save_docx(title, blocks, filename):
    doc = Document()
    doc.add_heading(title, level=1)
    for kind, content in blocks:
        if kind == "heading": doc.add_heading(content, level=2)
        elif kind == "paragraph": doc.add_paragraph(content)
        elif kind == "list": doc.add_paragraph(content, style="List Bullet")
        elif kind == "code":
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.font.name = "Courier New"
        elif kind == "table":
            if not content: continue
            table = doc.add_table(rows=len(content), cols=len(content[0]))
            table.style = 'Table Grid'
            for i, row_data in enumerate(content):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell_text
    doc.save(filename)

def main():
    url = input("Enter GFG URL: ").strip()
    fmt = input("Format (pdf/docx): ").strip().lower()
    try:
        title, article = fetch_article(url)
        blocks = parse_blocks(article)
        safe_name = "".join([c for c in title if c.isalnum() or c in (' ', '_')]).strip().replace(" ", "_")
        if fmt == "pdf":
            save_pdf(title, blocks, f"{safe_name}.pdf")
            print(f"Done: {safe_name}.pdf")
        else:
            save_docx(title, blocks, f"{safe_name}.docx")
            print(f"Done: {safe_name}.docx")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()